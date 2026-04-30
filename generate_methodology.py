"""
Generate IEEE-formatted Proposed Methodology section as a Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0


def add_heading_ieee(text, level=1):
    """Add IEEE-style heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.size = Pt(10)
        run.bold = True
        # IEEE section headings are roman numeral + centered
    elif level == 2:
        run.font.size = Pt(10)
        run.bold = False
        run.italic = True
    elif level == 3:
        run.font.size = Pt(10)
        run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_body(text, first_line_indent=True):
    """Add IEEE body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    return p


def add_figure_caption(text):
    """Add IEEE figure caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table_ieee(headers, rows, caption_text):
    """Add IEEE-style table with caption on top."""
    # Table caption above
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# III. PROPOSED METHODOLOGY
# ============================================================

add_heading_ieee("III. PROPOSED METHODOLOGY")

add_body(
    "This study proposes a noise-aware fine-tuning framework for BirdNET-based avian species "
    "classification, specifically designed to improve recognition accuracy on the Indian Bird Call "
    "dataset (IBC53) comprising 30 species. The proposed pipeline integrates domain-specific audio "
    "preprocessing with transfer learning on the BirdNET deep learning backbone. The framework "
    "consists of four principal stages: (1) audio segmentation, (2) energy-based noise detection, "
    "(3) noise-augmented dataset construction, and (4) BirdNET fine-tuning with multi-configuration "
    "evaluation. The overall architecture of the proposed methodology is depicted in Fig. 1."
)

# Insert architecture diagram
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(2)
run = p.add_run()
run.add_picture("/Users/adityas/Downloads/birdnet_pipeline_architecture (1).png", width=Inches(5.5))

add_figure_caption(
    "Fig. 1. Architecture of the proposed noise-aware BirdNET fine-tuning framework showing the "
    "four-stage pipeline: audio segmentation, energy-based noise detection, dataset construction "
    "with ESC-50 noise integration, and BirdNET fine-tuning with multi-experiment evaluation."
)

# ── A. Audio Segmentation ──
add_heading_ieee("A. Audio Segmentation", level=2)

add_body(
    "The first stage of the proposed pipeline addresses the challenge of processing variable-length "
    "field recordings into uniform input units suitable for deep learning-based classification. Raw "
    "audio files from the IBC53 dataset, totaling 1,252 recordings across 30 Indian bird species and "
    "an additional 443 unclassified (mystery) files, are segmented into fixed-duration chunks of 3 seconds "
    "each. This segment duration is chosen to align with BirdNET's native input specification, which "
    "expects 3-second audio windows sampled at 48 kHz."
)

add_body(
    "The segmentation process employs a non-overlapping sliding window approach. Each raw audio file is "
    "first resampled to 48 kHz to ensure uniform sampling rate across all recordings. The resampled "
    "waveform is then divided into consecutive 3-second segments. Segments shorter than 3 seconds at "
    "the tail end of a recording are zero-padded to maintain dimensional consistency. This stage "
    "transforms the original 1,252 recordings into approximately 4,854 uniform audio segments, "
    "establishing a standardized input representation for downstream processing."
)

add_body(
    "Let x(t) denote a raw audio signal of duration T seconds. The segmentation function S partitions "
    "x(t) into N segments, where N = ceil(T / 3). Each segment s_i is defined as:"
)

# Equation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("s_i = x(t) for t in [3(i-1), 3i],  i = 1, 2, ..., N          (1)")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

add_body(
    "If the length of the final segment is less than 3 seconds, zero-padding is applied to "
    "ensure all segments have a uniform length of 144,000 samples (3 seconds x 48,000 Hz)."
)

# ── B. Energy-Based Noise Detection ──
add_heading_ieee("B. Energy-Based Noise Detection", level=2)

add_body(
    "Field recordings inherently contain a mixture of target bird vocalizations, ambient environmental "
    "noise, and periods of silence. Training a classifier on unfiltered segments introduces label noise "
    "and degrades model performance. The second stage of the proposed pipeline addresses this by implementing "
    "an energy-based noise detection module that classifies each 3-second segment into one of three "
    "categories: bird vocalization, environmental noise, or silence."
)

add_body(
    "The classification is performed using three complementary acoustic features computed directly "
    "from the time-domain and frequency-domain representations of each segment:"
)

# Feature descriptions
add_body(
    "1) Root Mean Square (RMS) Energy: RMS energy quantifies the overall amplitude of the audio signal. "
    "Segments with RMS values below a threshold theta_RMS = 0.01 are classified as silence, as they "
    "contain insufficient acoustic energy to represent any meaningful vocalization or noise event. "
    "The RMS energy for a segment s_i of length L samples is computed as:",
    first_line_indent=True
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RMS(s_i) = sqrt((1/L) * sum(s_i[n]^2)),  n = 0, 1, ..., L-1          (2)")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

add_body(
    "2) Spectral Flatness: Spectral flatness measures the uniformity of the power spectrum. A high "
    "spectral flatness value indicates a flat, noise-like spectrum (resembling white noise), while a low "
    "value indicates a spectrum with distinct peaks characteristic of tonal signals such as bird calls. "
    "Spectral flatness is defined as the ratio of the geometric mean to the arithmetic mean of the "
    "power spectral density. Segments with spectral flatness exceeding theta_SF = 0.5 are flagged as "
    "candidate noise segments.",
    first_line_indent=True
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SF(s_i) = (prod(P(k))^(1/K)) / ((1/K) * sum(P(k))),  k = 0, ..., K-1          (3)")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

add_body(
    "where P(k) is the power spectral density at frequency bin k, and K is the total number of "
    "frequency bins.",
    first_line_indent=True
)

add_body(
    "3) Zero-Crossing Rate (ZCR): ZCR counts the number of times the signal crosses the zero amplitude "
    "axis per unit time. Environmental noise typically exhibits a high ZCR due to its aperiodic nature, "
    "whereas bird vocalizations tend to have more structured temporal patterns with lower ZCR. A segment "
    "is classified as noise if both its spectral flatness exceeds theta_SF and its ZCR exceeds theta_ZCR = 0.1.",
    first_line_indent=True
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ZCR(s_i) = (1/(2L)) * sum(|sign(s_i[n]) - sign(s_i[n-1])|),  n = 1, ..., L-1          (4)")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

add_body(
    "The combined decision rule for segment classification is formulated as follows: a segment is "
    "labeled as silence if RMS(s_i) < theta_RMS; as noise if SF(s_i) > theta_SF AND ZCR(s_i) > theta_ZCR; "
    "and as bird vocalization otherwise. Table I summarizes the threshold parameters used in the "
    "noise detection module."
)

# Table I
add_table_ieee(
    headers=["Parameter", "Symbol", "Value", "Purpose"],
    rows=[
        ["RMS Threshold", "theta_RMS", "0.01", "Silence detection"],
        ["Spectral Flatness Threshold", "theta_SF", "0.5", "Noise candidate detection"],
        ["Zero-Crossing Rate Threshold", "theta_ZCR", "0.1", "Noise confirmation"],
    ],
    caption_text="TABLE I. Threshold Parameters for Energy-Based Noise Detection"
)

# ── C. Dataset Construction with Noise Augmentation ──
add_heading_ieee("C. Dataset Construction with Noise Augmentation", level=2)

add_body(
    "The third stage constructs the training datasets required for BirdNET fine-tuning. A critical "
    "contribution of this work is the integration of dedicated noise samples into the training data "
    "to enable the model to learn noise suppression during inference. The dataset construction module "
    "produces multiple dataset variants to support the experimental configurations described in "
    "Section IV."
)

add_body(
    "The filtered bird vocalization segments from Stage 2 are organized into species-specific "
    "directories following BirdNET's expected training format, where each subdirectory is named using "
    "the scientific name of the species (e.g., Pnoepyga_pusilla, Pellorneum_ruficeps). The resulting "
    "bird-only dataset contains 6,924 audio files distributed across 30 species classes."
)

add_body(
    "To construct the noise-augmented dataset, environmental noise samples are extracted from the "
    "ESC-50 dataset [ref]. Seven noise categories relevant to field recording conditions are selected: "
    "rain, wind, thunderstorm, insects, water drops, crackling fire, and engine idling. A total of "
    "283 noise segments are extracted and placed into a dedicated noise directory. BirdNET treats this "
    "directory as a NON_EVENT_CLASS during training, meaning the noise samples do not produce a separate "
    "output label but instead train the model's internal representations to suppress detections on "
    "non-bird audio segments. The noise-augmented dataset contains 7,207 files (6,924 bird + 283 noise)."
)

add_body(
    "Additionally, few-shot subsets are constructed to evaluate data efficiency. For each of the 30 "
    "species and the noise class, random subsets of 10, 25, and 50 samples per class are drawn without "
    "replacement, resulting in three few-shot datasets of 583, 1,033, and 1,783 files respectively. "
    "Table II summarizes the dataset configurations."
)

# Table II
add_table_ieee(
    headers=["Dataset", "Bird Files", "Noise Files", "Total", "Classes"],
    rows=[
        ["Without Noise (Exp 1)", "6,924", "0", "6,924", "30"],
        ["With Noise (Exp 2)", "6,924", "283", "7,207", "30 + noise"],
        ["Few-Shot 10 (Exp 3a)", "300", "283", "583", "30 + noise"],
        ["Few-Shot 25 (Exp 3b)", "750", "283", "1,033", "30 + noise"],
        ["Few-Shot 50 (Exp 3c)", "1,500", "283", "1,783", "30 + noise"],
    ],
    caption_text="TABLE II. Dataset Configurations for Experimental Evaluation"
)

# ── D. BirdNET Fine-Tuning and Evaluation ──
add_heading_ieee("D. BirdNET Fine-Tuning and Evaluation", level=2)

add_body(
    "The fourth and final stage of the proposed pipeline performs transfer learning on the pre-trained "
    "BirdNET model using the constructed datasets. BirdNET is a deep convolutional neural network "
    "originally trained on over 6,000 bird species worldwide using the EfficientNet architecture as "
    "its backbone. The pre-trained model provides robust general-purpose audio feature extraction "
    "capabilities, which are adapted to the target domain of Indian bird species through fine-tuning."
)

add_body(
    "During fine-tuning, the pre-trained feature extraction layers of BirdNET are frozen, and a new "
    "classification head is trained on the target dataset. The classification head maps the extracted "
    "feature embeddings to the 30 species classes. The model is trained for 50 epochs with the training "
    "configuration managed through BirdNET's built-in training API. The best model checkpoint is selected "
    "based on the Area Under the Precision-Recall Curve (AUPRC), which is particularly suitable for "
    "imbalanced multi-class audio classification tasks."
)

add_body(
    "The fine-tuned model is exported in TensorFlow Lite (.tflite) format for efficient inference. "
    "Each model variant is approximately 25.1 MB in size, enabling deployment on resource-constrained "
    "devices for field applications."
)

add_body(
    "Four experimental configurations are evaluated to systematically assess the impact of noise-aware "
    "training and data volume:"
)

add_body(
    "1) Experiment 1 (Baseline Fine-Tuning): The model is fine-tuned on 6,924 bird segments across "
    "30 species without any noise class. This configuration serves as the baseline to quantify the "
    "improvement achieved by noise-aware training.",
    first_line_indent=True
)

add_body(
    "2) Experiment 2 (Noise-Aware Fine-Tuning): The model is fine-tuned on 7,207 files including "
    "the ESC-50 noise class. This is the key experiment of this study, designed to evaluate whether "
    "explicit noise training improves classification accuracy and confidence calibration.",
    first_line_indent=True
)

add_body(
    "3) Experiment 3 (Few-Shot Data Sensitivity): Three sub-experiments are conducted with 10, 25, "
    "and 50 samples per species to characterize the data efficiency curve of the fine-tuning process "
    "and identify the minimum viable training set size for acceptable performance.",
    first_line_indent=True
)

add_heading_ieee("E. Evaluation Metrics", level=2)

add_body(
    "All experiments are evaluated on the complete IBC53 test set comprising 1,252 audio files across "
    "30 species and 443 mystery (unclassified) files. The following metrics are computed to provide a "
    "comprehensive assessment of model performance:"
)

add_body(
    "1) Overall Accuracy: The ratio of correctly classified detections to total species detections, "
    "computed across all test files and all species.",
    first_line_indent=True
)

add_body(
    "2) Per-Species Accuracy: Individual classification accuracy for each of the 30 species, enabling "
    "identification of species-specific performance patterns and inter-species confusion.",
    first_line_indent=True
)

add_body(
    "3) Confidence Distribution: Statistical analysis of prediction confidence scores including mean, "
    "median, and the proportion of detections exceeding confidence thresholds of 0.5, 0.7, and 0.9. "
    "Higher confidence indicates better model calibration.",
    first_line_indent=True
)

add_body(
    "4) Confusion Analysis: Identification of the most frequent misclassification pairs to understand "
    "inter-species acoustic similarity and its impact on classification errors.",
    first_line_indent=True
)

add_body(
    "5) Training Metrics: AUPRC, AUROC, and training loss are monitored during the fine-tuning process "
    "to assess convergence behavior and model selection quality.",
    first_line_indent=True
)

add_body(
    "The evaluation framework enables direct comparison across all experimental configurations, "
    "isolating the individual contributions of noise-aware training and training data volume to "
    "classification performance on Indian bird species."
)

# ── Save ──
output_path = "/Users/adityas/Desktop/FALCON_DL/FALCON_DL/Proposed_Methodology_IEEE.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
